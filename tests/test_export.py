"""export_docx: only changed clauses move, everything else rides through."""

import io

from docx import Document

from app.export import export_docx


def _make_docx() -> bytes:
    doc = Document()
    doc.add_heading("SERVICES AGREEMENT", level=1)
    doc.add_paragraph("1. Payment", style="Heading 2")
    doc.add_paragraph("Payment is due within seventy five days of invoice.")
    doc.add_paragraph("2. Liability", style="Heading 2")
    doc.add_paragraph("Liability is unlimited for either party.")
    doc.add_paragraph("It survives termination of this agreement.")
    doc.add_paragraph("3. Notices", style="Heading 2")
    doc.add_paragraph("Notices must be sent by registered post.")
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _texts(data: bytes) -> list[str]:
    return [p.text for p in Document(io.BytesIO(data)).paragraphs if p.text]


def test_changed_clause_is_replaced_and_rest_untouched():
    original = _make_docx()
    clauses = [
        {"clause_id": "c01", "decision": "accepted",
         "text": "Payment is due within seventy five days of invoice.",
         "final_text": "Payment is due within thirty days of invoice."},
        {"clause_id": "c03", "decision": "rejected",
         "text": "Notices must be sent by registered post.",
         "final_text": "Notices must be sent by registered post."},
    ]
    corrected, unmatched = export_docx(original, clauses)
    texts = _texts(corrected)
    assert unmatched == []
    assert "Payment is due within thirty days of invoice." in texts
    assert "Payment is due within seventy five days of invoice." not in texts
    assert "Notices must be sent by registered post." in texts
    assert "SERVICES AGREEMENT" in texts  # untouched parts survive


def test_multi_paragraph_clause_collapses_to_one():
    original = _make_docx()
    clauses = [
        {"clause_id": "c02", "decision": "edited",
         "text": "Liability is unlimited for either party. It survives termination of this agreement.",
         "final_text": "Liability is capped at fees paid in the prior twelve months."},
    ]
    corrected, unmatched = export_docx(original, clauses)
    texts = _texts(corrected)
    assert unmatched == []
    assert "Liability is capped at fees paid in the prior twelve months." in texts
    assert "It survives termination of this agreement." not in texts


def test_heading_style_survives_the_rewrite():
    original = _make_docx()
    clauses = [
        {"clause_id": "c01", "decision": "edited",
         "text": "Payment is due within seventy five days of invoice.",
         "final_text": "Payment is due on receipt."},
    ]
    corrected, _ = export_docx(original, clauses)
    styles = {p.text: p.style.name for p in Document(io.BytesIO(corrected)).paragraphs}
    assert styles["1. Payment"] == "Heading 2"
    assert styles["Payment is due on receipt."] == "Normal"


def test_unlocatable_clause_is_reported_not_silently_dropped():
    original = _make_docx()
    clauses = [
        {"clause_id": "c09", "decision": "edited",
         "text": "wording the model reflowed and no longer matches the file",
         "final_text": "anything"},
    ]
    corrected, unmatched = export_docx(original, clauses)
    assert unmatched == ["c09"]
    assert "SERVICES AGREEMENT" in _texts(corrected)  # file still exports
