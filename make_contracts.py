# make_contracts.py: builds the Papyrus synthetic contract set.
# Templates in, contracts + ground-truth manifests out. No model calls. $0.
import argparse
import json
from pathlib import Path

from docx import Document
from fpdf import FPDF
from fpdf.enums import XPos, YPos

ROOT = Path(__file__).resolve().parent
TEMPLATE_DIR = ROOT / "data" / "templates"
CONTRACT_DIR = ROOT / "data" / "contracts"
MANIFEST_DIR = ROOT / "data" / "manifests"
GOLDEN_DIR = ROOT / "tests" / "golden"

FIRM = "Aldergate Consulting Pte Ltd"
DATE_LINE = "Dated 1 July 2026"

TITLES = {
    "nda": "Mutual Non-Disclosure Agreement",
    "msa": "Master Services Agreement",
    "sow": "Statement of Work",
    "vendor": "Vendor Agreement",
    "employment": "Employment Agreement",
}

INSPECTORS = {"compliance", "risk", "template", "financial"}
SEVERITIES = {"high", "medium", "low"}
def _d(dev_id, ctype, clause_type, heading, severity, inspector, note, text):
    return {"id": dev_id, "contract_type": ctype, "clause_type": clause_type,
            "heading": heading, "severity": severity, "inspector": inspector,
            "note": note, "deviated_text": text}


NDA_DEVIATIONS = [
    _d("nda-scope-oneway", "nda", "scope", "Confidential Information", "high", "risk",
       "One-way definition: only the counterparty's information is protected, ours is not.",
       "Confidential Information means any non-public information disclosed by [COUNTERPARTY] to the "
       "other party. Information disclosed by the other party to [COUNTERPARTY] is not Confidential "
       "Information and carries no obligations under this Agreement."),
    _d("nda-conf-perpetual", "nda", "term_termination", "Term and Survival", "medium", "template",
       "Survival changed from the standard three years to perpetual.",
       "This Agreement runs for five (5) years from the date of signature. The confidentiality "
       "obligations survive indefinitely, without any limit in time."),
    _d("nda-dp-overseas", "nda", "data_protection", "Data Protection", "high", "compliance",
       "Personal data may go to any affiliate in any country, with no consent or safeguards.",
       "The receiving party may transfer any personal data received under this Agreement to its "
       "affiliates and service providers in any country, without further consent and without "
       "additional safeguards."),
    _d("nda-gov-offshore", "nda", "governing_law", "Governing Law", "medium", "compliance",
       "Governing law moved from Singapore to the Cayman Islands.",
       "This Agreement is governed by the laws of the Cayman Islands, and any dispute must be "
       "referred to confidential arbitration seated in George Town."),
    _d("nda-compelled-nonotice", "nda", "confidentiality", "Compelled Disclosure", "low", "compliance",
       "Compelled disclosure allowed without notifying us first.",
       "The receiving party may disclose Confidential Information whenever it believes disclosure is "
       "required by law or requested by any authority, without notice to the disclosing party and "
       "without limiting the scope of what is disclosed."),
    _d("nda-conf-carveout", "nda", "confidentiality", "Obligations of Confidentiality", "high", "compliance",
       "Our information may be passed to any contractor or affiliate with no confidentiality undertakings.",
       "The receiving party will use reasonable efforts to keep Confidential Information confidential, "
       "and may share it with any contractor, consultant, or affiliate it chooses without requiring "
       "any confidentiality undertaking from them."),
    _d("nda-ip-grab", "nda", "ip", "No Licence", "high", "risk",
       "Receiving our information grants a perpetual licence to use it for any purpose.",
       "Disclosure of Confidential Information under this Agreement grants the receiving party a "
       "perpetual, irrevocable, royalty-free licence to use that information for any purpose, "
       "including in its own products and services."),
    _d("nda-liq-damages", "nda", "liability", "Remedies", "high", "financial",
       "A fixed S$500,000 penalty per disclosure replaces the standard remedies clause.",
       "The receiving party must pay the disclosing party liquidated damages of five hundred thousand "
       "Singapore dollars (S$500,000) for each individual disclosure made in breach of this Agreement, "
       "in addition to any other remedy available to the disclosing party."),
    _d("nda-return-keep", "nda", "confidentiality", "Return or Destruction", "medium", "template",
       "The counterparty may keep all our information forever instead of returning or destroying it.",
       "The receiving party may retain copies of all Confidential Information for its records and "
       "internal use, and has no obligation to return or destroy any materials when this Agreement ends."),
    _d("nda-notices-fax", "nda", "notices", "Notices", "low", "template",
       "Notices are valid only by fax, a channel the firm does not use.",
       "Formal notices under this Agreement are valid only if sent by facsimile to the number stated "
       "in the signature block, and take effect on transmission."),
]
MSA_DEVIATIONS = [
    _d("msa-pay-net120", "msa", "payment", "Fees and Payment", "high", "financial",
       "Payment stretched from net-30 to net-120, with instant suspension if late.",
       "The Client pays each invoice within one hundred and twenty (120) days of the invoice date. "
       "The Provider may suspend all services immediately, without notice, if any invoice remains "
       "unpaid after that period."),
    _d("msa-liability-unlimited", "msa", "liability", "Limitation of Liability", "high", "risk",
       "The Client's liability is unlimited while the Provider's is capped at ten percent of three months of fees.",
       "The Client's liability under this Agreement is unlimited. The Provider's total liability for "
       "all claims is capped at ten percent (10%) of the fees paid in the three (3) months before the "
       "claim, and the Provider is not liable for indirect or consequential loss."),
    _d("msa-term-autorenew", "msa", "term_termination", "Term and Termination", "medium", "risk",
       "Auto-renewal trap: 180 days notice to escape, and only the Provider may terminate for convenience.",
       "This Agreement renews automatically for successive twelve (12) month terms unless the Client "
       "gives written notice at least one hundred and eighty (180) days before the renewal date. The "
       "Provider may terminate for convenience at any time on notice; the Client may otherwise "
       "terminate only for unremedied material breach."),
    _d("msa-ip-provider-owns", "msa", "ip", "Intellectual Property", "high", "template",
       "Deliverables stay the Provider's property; the Client only gets a revocable internal-use licence.",
       "All deliverables and work product remain the exclusive property of the Provider. The Client "
       "receives a non-exclusive, non-transferable licence to use the deliverables for internal "
       "purposes only, which the Provider may revoke on termination of this Agreement."),
    _d("msa-dp-nobreach", "msa", "data_protection", "Data Protection", "high", "compliance",
       "No PDPA commitment and no duty to notify the Client of a data breach.",
       "The Provider will handle personal data in accordance with its own internal privacy practices. "
       "The Provider has no obligation to notify the Client of any data incident, and may engage "
       "sub-processors at its discretion."),
    _d("msa-pay-interest8", "msa", "payment", "Fees and Payment", "medium", "financial",
       "Late interest raised to eight percent per month, compounded, running from the invoice date.",
       "The Client pays undisputed invoices within thirty (30) days of receipt. Late amounts accrue "
       "interest at eight percent (8%) per month, compounded monthly, running from the invoice date "
       "rather than the due date."),
    _d("msa-restraint-noncompete", "msa", "restraint", "Non-Solicitation", "high", "compliance",
       "Two-year non-compete blocking the Client from similar providers and from hiring anyone connected to the Provider.",
       "For two (2) years after this Agreement ends, the Client will not engage any other provider of "
       "the same or similar services, and will not employ or solicit any person who was employed or "
       "engaged by the Provider, in any capacity."),
    _d("msa-conf-oneway", "msa", "confidentiality", "Confidentiality", "high", "risk",
       "One-way confidentiality: the Client's information gets no protection.",
       "The Client must keep the Provider's Confidential Information secret and use it only to perform "
       "this Agreement. The Provider accepts no confidentiality obligation in respect of information "
       "received from the Client."),
    _d("msa-warranty-asis", "msa", "other", "Service Warranty", "medium", "template",
       "The skill-and-care warranty is replaced by an as-is disclaimer.",
       "The services are provided on an as is and as available basis, with all faults. The Provider "
       "disclaims all warranties and conditions, express or implied, to the maximum extent permitted "
       "by law."),
    _d("msa-term-oneway", "msa", "term_termination", "Term and Termination", "medium", "risk",
       "Exit rights one-sided: the Provider may leave on 7 days notice while the Client is locked in.",
       "The Provider may terminate this Agreement at any time on seven (7) days written notice. The "
       "Client may terminate only at the end of the initial term, on not less than ninety (90) days "
       "prior written notice, and otherwise remains bound for the full term."),
]
SOW_DEVIATIONS = [
    _d("sow-pay-upfront", "sow", "payment", "Fees and Payment", "high", "financial",
       "All fees due on signature, non-refundable, before any work starts.",
       "The Client pays one hundred percent (100%) of the total fees stated in the Fee Schedule on "
       "signature of this Statement of Work, before any work begins. All fees are non-refundable in "
       "all circumstances."),
    _d("sow-scope-unilateral", "sow", "scope", "Deliverables", "high", "risk",
       "The Provider may change scope and timeline at will and bill extra work at its own rates.",
       "The Provider may vary the scope, deliverables, and timeline at its reasonable discretion as "
       "the work proceeds. Any additional work the Provider performs is chargeable at the Provider's "
       "then-current rates without a change order."),
    _d("sow-acceptance-deemed", "sow", "other", "Acceptance", "medium", "template",
       "Deliverables deemed accepted on delivery, with no right to reject.",
       "Each deliverable is deemed accepted on delivery. The Client waives any right to reject, or to "
       "require correction of, a deliverable after it has been delivered."),
    _d("sow-ip-reuse", "sow", "ip", "Intellectual Property", "high", "compliance",
       "The Provider keeps deliverables until year end and may reuse work containing Client materials for other customers.",
       "The Provider retains ownership of all deliverables until the end of the calendar year in which "
       "they are delivered, and may reuse any deliverable, including Client materials incorporated in "
       "it, for other customers."),
    _d("sow-cancel-fee", "sow", "term_termination", "Term and Cancellation", "medium", "financial",
       "Cancelling costs the full remaining fees plus fifteen percent.",
       "If the Client cancels any part of the work for any reason, the Client must pay a cancellation "
       "fee equal to the full remaining fees under this Statement of Work plus fifteen percent (15%) "
       "of the total contract value."),
]
VENDOR_DEVIATIONS = [
    _d("ven-price-anytime", "vendor", "payment", "Prices and Payment", "high", "financial",
       "Prices changeable at any time with no notice, applying even to confirmed orders.",
       "The Vendor may revise its prices at any time without notice. Revised prices apply to all "
       "orders not yet delivered, including orders already confirmed by the Client."),
    _d("ven-delivery-risk", "vendor", "other", "Delivery and Risk", "medium", "template",
       "Risk moves to the Client when goods leave the warehouse, so the Client pays for transit damage.",
       "Risk in goods passes to the Client as soon as the goods leave the Vendor's premises. The "
       "Client bears the cost of any loss of or damage to goods in transit, and the Vendor has no "
       "liability for delivery failures once goods are dispatched."),
    _d("ven-indemnity-flip", "vendor", "indemnity", "Indemnity", "high", "risk",
       "Indemnity reversed: the Client covers the Vendor's losses, even those caused by the Vendor's own negligence.",
       "The Client indemnifies the Vendor against all claims, losses, damages, and costs arising in "
       "connection with this Agreement, including claims arising from the Vendor's own negligence "
       "or breach."),
    _d("ven-dp-subprocessors", "vendor", "data_protection", "Data Protection", "medium", "compliance",
       "Any sub-processor in any country may handle personal data, with no notice or consent.",
       "The Vendor may appoint any sub-processor to handle personal data under this Agreement without "
       "notice to the Client, and personal data may be stored or processed in any jurisdiction the "
       "Vendor selects."),
    _d("ven-notices-post-only", "vendor", "notices", "Notices", "low", "template",
       "Email notices are void; registered post is the only valid channel.",
       "Notices under this Agreement are valid only if sent by registered post to the Vendor's "
       "registered office. Notices sent by email or any other means are of no effect."),
    _d("ven-pay-net90", "vendor", "payment", "Prices and Payment", "medium", "financial",
       "Net-90 payment plus a two percent fee whenever the Client queries an invoice.",
       "The Client pays invoices within ninety (90) days of the invoice date. The Vendor may add an "
       "administration fee of two percent (2%) to any invoice that the Client disputes or queries."),
    _d("ven-inspect-48h", "vendor", "other", "Quality and Inspection", "medium", "template",
       "Inspection window cut from ten business days to 48 hours, then all claims are waived.",
       "The Client must notify the Vendor of any defect or shortfall within forty-eight (48) hours of "
       "delivery. After that window the goods are deemed accepted and the Client waives all claims "
       "relating to them."),
    _d("ven-term-lock", "vendor", "term_termination", "Term and Termination", "medium", "compliance",
       "Five-year lock-in with no exit for the Client, while the Vendor may leave at any time.",
       "The initial term of this Agreement is five (5) years. The Client may not terminate for "
       "convenience during the initial term. The Vendor may terminate this Agreement at any time on "
       "thirty (30) days written notice."),
    _d("ven-liability-strike", "vendor", "liability", "Limitation of Liability", "high", "risk",
       "The Vendor's liability is excluded entirely, even for defective goods, while the Client's is unlimited.",
       "The Vendor excludes all liability under this Agreement to the fullest extent permitted by "
       "law, including liability for defective goods and late delivery. The Client's liability to "
       "the Vendor is unlimited."),
    _d("ven-supply-asis", "vendor", "scope", "Supply", "high", "template",
       "Goods supplied as seen, with no promise of matching any specification.",
       "The Vendor will supply the goods and services listed in the Order Schedule on an as seen "
       "basis. The Vendor gives no assurance that goods will match any specification, sample, or "
       "description, and substitutions are at the Vendor's discretion."),
]
EMPLOYMENT_DEVIATIONS = [
    _d("emp-restraint-2y", "employment", "restraint", "Non-Solicitation", "high", "compliance",
       "Worldwide two-year non-compete plus a total contact ban.",
       "For twenty-four (24) months after the employment ends, the Employee will not work for, advise, "
       "or hold any interest in any competing business anywhere in the world, and will not contact any "
       "client or employee of the Employer for any purpose."),
    _d("emp-bonus-guaranteed", "employment", "payment", "Salary and Benefits", "medium", "financial",
       "A guaranteed bonus of six months of salary regardless of performance.",
       "In addition to the monthly salary stated in Schedule A, the Employee receives a guaranteed "
       "annual bonus equal to six (6) months of salary, payable regardless of the Employee's "
       "performance or the Employer's results."),
    _d("emp-probation-12m", "employment", "other", "Probation", "medium", "compliance",
       "Probation stretched to twelve months with dismissal allowed with no notice at all.",
       "The probationary period is twelve (12) months. During probation the Employer may end the "
       "employment at any time without notice, without payment in lieu, and without giving a reason."),
    _d("emp-ip-moonlight", "employment", "ip", "Intellectual Property", "medium", "template",
       "The Employer claims IP in everything created during employment, even private work off the clock.",
       "All intellectual property the Employee creates at any time during the period of employment "
       "belongs to the Employer, including work created outside working hours, without Employer "
       "resources, and unrelated to the Employer's business."),
    _d("emp-term-oneway", "employment", "term_termination", "Termination", "medium", "risk",
       "The Employer may dismiss without notice while the Employee owes six months notice to resign.",
       "The Employer may end the employment at any time without notice or payment in lieu. The "
       "Employee must give six (6) months written notice to resign, and the Employer may hold the "
       "Employee to garden leave for that full period."),
]


# --- coverage set: the clause types the first thirteen never exercised --------
# parties and boilerplate had ZERO planted defects, governing_law and indemnity
# had one each. A clause type with no defect is a clause type the inspectors are
# never scored on, so those four were blind spots in the labelled truth rather
# than in the code.
COVERAGE_DEVIATIONS = [
    _d("cov-parties-undisclosed", "msa", "parties", "Parties", "medium", "compliance",
       "The counterparty may substitute any group company as the contracting entity without telling us.",
       "This Agreement is between the Client and the Supplier, provided that the Supplier may at any time and "
       "without notice novate or assign this Agreement to any entity within its group, and the Client's consent "
       "shall not be required and need not be sought."),
    _d("cov-boiler-unilateral", "msa", "boilerplate", "General", "high", "risk",
       "The counterparty can change the terms of the contract on its own by posting an update.",
       "The Supplier may amend these terms at any time by publishing an updated version on its website, and such "
       "amendment shall take effect immediately and bind the Client whether or not the Client has read it. No "
       "signature or countersignature of the Client is required."),
    _d("cov-indemnity-uncapped", "msa", "indemnity", "Indemnity", "high", "risk",
       "We indemnify them for everything, without limit, even when the fault is theirs.",
       "The Client shall indemnify, defend and hold harmless the Supplier against all claims, losses, damages and "
       "expenses of any kind whatsoever, without cap or limitation, including those arising from the Supplier's own "
       "negligence, wilful misconduct or breach of this Agreement."),
    _d("cov-gov-arbitration", "msa", "governing_law", "Governing Law", "medium", "compliance",
       "Any dispute has to be arbitrated abroad, in a seat and language that suit only them.",
       "This Agreement is governed by the laws of the British Virgin Islands. Any dispute shall be referred to "
       "arbitration seated in Tortola conducted in a language nominated by the Supplier, and the Client irrevocably "
       "waives any right to bring proceedings in any other forum."),
    _d("cov-parties-authority", "vendor", "parties", "Parties", "low", "template",
       "Nobody has confirmed the signatory can actually bind the counterparty.",
       "This Agreement is entered into by the Supplier acting through any person who presents themselves as "
       "authorised, and the Supplier gives no warranty that any such person holds authority to bind it."),
    _d("cov-boiler-noassign", "vendor", "boilerplate", "General", "medium", "template",
       "We cannot assign the contract even in a group reorganisation, but they can assign it freely.",
       "The Customer shall not assign, novate or otherwise transfer this Agreement or any part of it under any "
       "circumstances, including intra-group reorganisation. The Supplier may assign this Agreement freely and "
       "without notice."),
    _d("cov-indemnity-noticebar", "vendor", "indemnity", "Indemnity", "medium", "compliance",
       "Any indemnity claim is barred unless we notice it within five days.",
       "No indemnity shall be payable unless the Customer gives written notice of the claim within five (5) days of "
       "the event giving rise to it, time being of the essence, failing which the claim is absolutely barred."),
]

DEVIATIONS = (NDA_DEVIATIONS + MSA_DEVIATIONS + SOW_DEVIATIONS
              + VENDOR_DEVIATIONS + EMPLOYMENT_DEVIATIONS + COVERAGE_DEVIATIONS)
CONTRACT_SET = [
    {"file": "nda_halcyon.docx", "contract_type": "nda",
     "counterparty": "Halcyon Logistics Pte Ltd", "deviations": []},
    {"file": "nda_meridian.docx", "contract_type": "nda",
     "counterparty": "Meridian Data Systems Pte Ltd",
     "deviations": ["nda-scope-oneway", "nda-conf-perpetual", "nda-dp-overseas",
                    "nda-gov-offshore", "nda-compelled-nonotice"]},
    {"file": "nda_kestrel.docx", "contract_type": "nda",
     "counterparty": "Kestrel Marine Services Ltd",
     "deviations": ["nda-conf-carveout", "nda-ip-grab", "nda-liq-damages",
                    "nda-return-keep", "nda-notices-fax"]},
    {"file": "msa_tembusu.docx", "contract_type": "msa",
     "counterparty": "Tembusu Analytics Pte Ltd", "deviations": []},
    {"file": "msa_novabright.docx", "contract_type": "msa",
     "counterparty": "Novabright Media LLP",
     "deviations": ["msa-pay-net120", "msa-liability-unlimited", "msa-term-autorenew",
                    "msa-ip-provider-owns", "msa-dp-nobreach"]},
    {"file": "msa_cobalt.docx", "contract_type": "msa",
     "counterparty": "Cobalt Harbour Ventures Pte Ltd",
     "deviations": ["msa-pay-interest8", "msa-restraint-noncompete", "msa-conf-oneway",
                    "msa-warranty-asis", "msa-term-oneway"],
     "omit": ["Data Protection"]},
    {"file": "sow_straitsforge.docx", "contract_type": "sow",
     "counterparty": "Straits Forge Engineering Pte Ltd", "deviations": []},
    {"file": "sow_windmoor.docx", "contract_type": "sow",
     "counterparty": "Windmoor Freight Solutions Pte Ltd",
     "deviations": ["sow-pay-upfront", "sow-scope-unilateral", "sow-acceptance-deemed",
                    "sow-ip-reuse", "sow-cancel-fee"]},
    {"file": "vendor_aurelia.docx", "contract_type": "vendor",
     "counterparty": "Aurelia Health Group Pte Ltd", "deviations": []},
    {"file": "vendor_larkspur.docx", "contract_type": "vendor",
     "counterparty": "Larkspur Talent Partners LLP",
     "deviations": ["ven-price-anytime", "ven-delivery-risk", "ven-indemnity-flip",
                    "ven-dp-subprocessors", "ven-notices-post-only"],
     "omit": ["Limitation of Liability"]},
    {"file": "vendor_brightquay.pdf", "contract_type": "vendor",
     "counterparty": "Brightquay Facilities Management Pte Ltd",
     "deviations": ["ven-pay-net90", "ven-inspect-48h", "ven-term-lock",
                    "ven-liability-strike", "ven-supply-asis"]},
    {"file": "employment_raghavan.docx", "contract_type": "employment",
     "counterparty": "Priya Raghavan", "deviations": []},
    {"file": "employment_ong.docx", "contract_type": "employment",
     "counterparty": "Marcus Ong",
     "deviations": ["emp-restraint-2y", "emp-bonus-guaranteed", "emp-probation-12m",
                    "emp-ip-moonlight", "emp-term-oneway"],
     "omit": ["Personal Data"]},
    # the coverage pair: added so parties, boilerplate, governing_law and indemnity
    # each carry planted defects. The original thirteen are untouched, so their
    # manifests are byte-identical and their per-contract bench scores still stand.
    {"file": "msa_ferrow.docx", "contract_type": "msa",
     "counterparty": "Ferrow Industrial Group Pte Ltd",
     "deviations": ["cov-parties-undisclosed", "cov-boiler-unilateral",
                    "cov-indemnity-uncapped", "cov-gov-arbitration"]},
    {"file": "vendor_calderwood.docx", "contract_type": "vendor",
     "counterparty": "Calderwood Supply Co Pte Ltd",
     "deviations": ["cov-parties-authority", "cov-boiler-noassign",
                    "cov-indemnity-noticebar"],
     "omit": ["Governing Law"]},
]
def load_clauses(contract_type: str) -> list[dict]:
    path = TEMPLATE_DIR / f"{contract_type}.json"
    return json.loads(path.read_text())["clauses"]


def build_body(spec: dict) -> tuple[list[tuple[str, str, str]], list[dict]]:
    """Returns (rows, planted): rows are (number, heading, text) ready to print."""
    catalogue = {d["id"]: d for d in DEVIATIONS}
    chosen = []
    for dev_id in spec["deviations"]:
        dev = catalogue[dev_id]
        assert dev["contract_type"] == spec["contract_type"], f"{dev_id} used on wrong contract type"
        chosen.append(dev)
    by_heading = {d["heading"]: d for d in chosen}
    assert len(by_heading) == len(chosen), f"{spec['file']}: two deviations aim at the same clause"
    omit = set(spec.get("omit", []))
    assert not (omit & set(by_heading)), f"{spec['file']}: cannot omit and deviate the same clause"

    rows, planted, n = [], [], 0
    for clause in load_clauses(spec["contract_type"]):
        if clause["heading"] in omit:
            assert clause["required"], f"{spec['file']}: omitting optional clause {clause['heading']} proves nothing"
            planted.append({"missing": True, "clause_type": clause["clause_type"],
                            "heading": clause["heading"]})
            omit.remove(clause["heading"])
            continue
        n += 1
        dev = by_heading.get(clause["heading"])
        text = dev["deviated_text"] if dev else clause["standard_text"]
        text = text.replace("the counterparty named in the signature block", spec["counterparty"])
        text = text.replace("the employee named in the signature block", spec["counterparty"])
        text = text.replace("[COUNTERPARTY]", spec["counterparty"])
        rows.append((str(n), clause["heading"], text))
        if dev:
            planted.append({"id": dev["id"], "number": str(n),
                            "clause_type": dev["clause_type"], "heading": dev["heading"],
                            "severity": dev["severity"], "inspector": dev["inspector"],
                            "note": dev["note"]})
    assert not omit, f"{spec['file']}: omit headings not found in template: {sorted(omit)}"
    applied = {p["heading"] for p in planted}
    unapplied = [d["id"] for d in chosen if d["heading"] not in applied]
    assert not unapplied, f"{spec['file']}: deviation headings not found in template: {unapplied}"
    return rows, planted
def write_docx(path: Path, spec: dict, rows: list) -> None:
    doc = Document()
    doc.add_heading(TITLES[spec["contract_type"]], level=1)
    doc.add_paragraph(DATE_LINE)
    doc.add_paragraph(f"Between {FIRM} and {spec['counterparty']}.")
    for number, heading, text in rows:
        doc.add_heading(f"{number}. {heading}", level=2)
        doc.add_paragraph(text)
    doc.add_paragraph(f"Signed: {FIRM}")
    doc.add_paragraph(f"Signed: {spec['counterparty']}")
    doc.save(path)


def write_pdf(path: Path, spec: dict, rows: list) -> None:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    def para(text, style="", size=11, h=6):
        pdf.set_font("Helvetica", style, size)
        pdf.multi_cell(0, h, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    para(TITLES[spec["contract_type"]], style="B", size=16, h=9)
    para(DATE_LINE)
    para(f"Between {FIRM} and {spec['counterparty']}.")
    for number, heading, text in rows:
        pdf.ln(3)
        para(f"{number}. {heading}", style="B", size=12, h=7)
        para(text)
    pdf.ln(6)
    para(f"Signed: {FIRM}")
    para(f"Signed: {spec['counterparty']}")
    pdf.output(str(path))


def write_manifest(spec: dict, planted: list) -> None:
    manifest = {"file": spec["file"], "contract_type": spec["contract_type"], "planted": planted}
    out = MANIFEST_DIR / (Path(spec["file"]).stem + ".manifest.json")
    out.write_text(json.dumps(manifest, indent=2) + "\n")
def write_tiny() -> None:
    GOLDEN_DIR.mkdir(parents=True, exist_ok=True)
    keep = ["Parties", "Obligations of Confidentiality", "Governing Law"]
    clauses = [c for c in load_clauses("nda") if c["heading"] in keep]
    assert len(clauses) == 3, f"expected 3 tiny clauses, found {len(clauses)}"
    doc = Document()
    doc.add_heading("Tiny Test Agreement", level=1)
    for n, clause in enumerate(clauses, start=1):
        text = clause["standard_text"].replace(
            "the counterparty named in the signature block", "Tembusu Analytics Pte Ltd")
        doc.add_heading(f"{n}. {clause['heading']}", level=2)
        doc.add_paragraph(text)
    doc.save(GOLDEN_DIR / "tiny.docx")
    print(f"wrote {GOLDEN_DIR / 'tiny.docx'} (3 clauses)")


def check_catalogue() -> None:
    ids = [d["id"] for d in DEVIATIONS]
    assert len(set(ids)) == len(ids), "duplicate deviation ids"
    for d in DEVIATIONS:
        assert d["inspector"] in INSPECTORS, f"{d['id']}: bad inspector {d['inspector']}"
        assert d["severity"] in SEVERITIES, f"{d['id']}: bad severity {d['severity']}"


def main() -> None:
    parser = argparse.ArgumentParser(description="generate the Papyrus test contracts")
    parser.add_argument("--tiny", action="store_true", help="only write tests/golden/tiny.docx")
    args = parser.parse_args()
    if args.tiny:
        write_tiny()
        return

    check_catalogue()
    CONTRACT_DIR.mkdir(parents=True, exist_ok=True)
    MANIFEST_DIR.mkdir(parents=True, exist_ok=True)
    total_planted = total_missing = 0
    for spec in CONTRACT_SET:
        n_dev = len(spec["deviations"])
        assert n_dev == 0 or 3 <= n_dev <= 5, f"{spec['file']}: {n_dev} deviations (want 0 or 3 to 5)"
        rows, planted = build_body(spec)
        path = CONTRACT_DIR / spec["file"]
        if spec["file"].endswith(".pdf"):
            write_pdf(path, spec, rows)
        else:
            write_docx(path, spec, rows)
        write_manifest(spec, planted)
        n_missing = sum(1 for p in planted if p.get("missing"))
        n_replaced = len(planted) - n_missing
        total_planted += n_replaced
        total_missing += n_missing
        tag = "clean" if not planted else f"{n_replaced} planted"
        if n_missing:
            tag += f" + {n_missing} missing"
        print(f"wrote {spec['file']:26} {spec['contract_type']:11} {tag}")
    print(f"\n{len(CONTRACT_SET)} contracts, {total_planted} planted deviations, "
          f"{total_missing} omitted required clauses, manifests in {MANIFEST_DIR}")


if __name__ == "__main__":
    main()
